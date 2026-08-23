# -*- coding: utf-8 -*-
"""Arma Proyecto_v6.docx reemplazando las secciones de diseño del Proyecto_v5.

Trabaja sobre el .docx original: conserva portada, indice, estilos, encabezados y
las secciones que no cambian. De cada seccion que si cambia borra el contenido que
hay entre su titulo y el titulo siguiente, y escribe el que generan los modulos de
docs/: los casos de uso, el modelo de datos, el diccionario de clases y las imagenes
de docs/diagramas.
"""
import copy
import os
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm

AQUI = os.path.dirname(os.path.abspath(__file__))
RAIZ = os.path.dirname(AQUI)
DIAGRAMAS = os.path.join(AQUI, 'diagramas')
sys.path.insert(0, AQUI)
sys.path.insert(0, DIAGRAMAS)

import modelo_datos  # noqa: E402
import diccionario_clases  # noqa: E402
from render_casos_de_uso import cargar, lineas  # noqa: E402

ENTRADA = os.path.join(RAIZ, 'Proyecto_v5.docx')
SALIDA = os.path.join(RAIZ, 'Proyecto_v6.docx')

ANCHO_MAXIMO_CM = 16.0


# --------------------------------------------------------------------------- utilidades

class Documento:
    """Cursor sobre el cuerpo del documento: borra secciones y escribe detras."""

    def __init__(self, ruta):
        self.doc = Document(ruta)
        self.cuerpo = self.doc.element.body
        self.cursor = None
        self._modelo_parrafo = None
        self._modelo_tabla = None
        # El indice del comienzo repite los titulos de todas las secciones. Sin este
        # corte, buscar un titulo devuelve su entrada del indice y el contenido nuevo
        # termina escrito ahi.
        self.inicio = self._fin_del_indice()

    def _fin_del_indice(self):
        for i, elemento in enumerate(self._elementos()):
            if not elemento.tag.endswith('}p'):
                continue
            pPr = elemento.find(qn('w:pPr'))
            estilo = pPr.find(qn('w:pStyle')) if pPr is not None else None
            if estilo is not None and estilo.get(qn('w:val')) == 'Heading1':
                if self._texto(elemento).startswith('2.'):
                    return i
        raise LookupError('no se encontro el comienzo del cuerpo del documento')

    # -- localizacion ------------------------------------------------------

    def _elementos(self):
        return list(self.cuerpo.iterchildren())

    def _texto(self, elemento):
        if not elemento.tag.endswith('}p'):
            return ''
        return ''.join(n.text or '' for n in elemento.iter(qn('w:t'))).strip()

    def buscar(self, prefijo, desde=None):
        desde = self.inicio if desde is None else desde
        for i, elemento in enumerate(self._elementos()):
            if i >= desde and self._texto(elemento).startswith(prefijo):
                return i
        raise LookupError('no se encontro: ' + prefijo)

    # -- edicion -----------------------------------------------------------

    def vaciar(self, titulo, hasta):
        """Borra lo que hay entre un titulo y el siguiente, y deja el cursor ahi."""
        elementos = self._elementos()
        i = self.buscar(titulo)
        j = self.buscar(hasta, i + 1)
        for elemento in elementos[i + 1:j]:
            self.cuerpo.remove(elemento)
        self.cursor = elementos[i]
        return self.cursor

    def modelo_parrafo(self):
        if self._modelo_parrafo is None:
            for elemento in self._elementos():
                if elemento.tag.endswith('}p') and self._texto(elemento) and \
                        elemento.find(qn('w:pPr')) is not None:
                    estilo = elemento.find(qn('w:pPr')).find(qn('w:pStyle'))
                    if estilo is None:
                        self._modelo_parrafo = elemento
                        break
        return self._modelo_parrafo

    def parrafo(self, texto='', estilo=None, negrita=False):
        nuevo = copy.deepcopy(self.modelo_parrafo())
        for hijo in list(nuevo):
            if not hijo.tag.endswith('}pPr'):
                nuevo.remove(hijo)

        pPr = nuevo.find(qn('w:pPr'))
        if estilo:
            for viejo in pPr.findall(qn('w:pStyle')):
                pPr.remove(viejo)
            pStyle = pPr.makeelement(qn('w:pStyle'), {qn('w:val'): estilo})
            pPr.insert(0, pStyle)

        if texto:
            run = nuevo.makeelement(qn('w:r'), {})
            if negrita:
                rPr = run.makeelement(qn('w:rPr'), {})
                rPr.append(rPr.makeelement(qn('w:b'), {}))
                run.append(rPr)
            t = run.makeelement(qn('w:t'), {})
            t.set(qn('xml:space'), 'preserve')
            t.text = texto
            run.append(t)
            nuevo.append(run)

        self.cursor.addnext(nuevo)
        self.cursor = nuevo
        return nuevo

    def titulo(self, texto, nivel=4):
        return self.parrafo(texto, estilo='Heading%d' % nivel)

    def modelo_tabla(self):
        if self._modelo_tabla is None:
            for elemento in self._elementos():
                if elemento.tag.endswith('}tbl'):
                    self._modelo_tabla = elemento
                    break
        return self._modelo_tabla

    def tabla(self, encabezados, filas):
        """Inserta una tabla con el mismo borde que las del documento original."""
        modelo = self.modelo_tabla()
        nueva = copy.deepcopy(modelo)
        for hijo in list(nueva):
            if hijo.tag.endswith('}tr'):
                nueva.remove(hijo)
        modelo_fila = copy.deepcopy(modelo.findall(qn('w:tr'))[0])

        # La grilla tiene que tener tantas columnas como la tabla nueva.
        grid = nueva.find(qn('w:tblGrid'))
        if grid is not None:
            for col in list(grid):
                grid.remove(col)
            ancho = int(9000 / len(encabezados))
            for _ in encabezados:
                grid.append(grid.makeelement(qn('w:gridCol'), {qn('w:w'): str(ancho)}))

        def fila(valores, negrita=False):
            tr = copy.deepcopy(modelo_fila)
            celdas = tr.findall(qn('w:tc'))
            base = copy.deepcopy(celdas[0])
            for celda in celdas:
                tr.remove(celda)
            for valor in valores:
                tc = copy.deepcopy(base)
                for p in tc.findall(qn('w:p')):
                    tc.remove(p)
                p = tc.makeelement(qn('w:p'), {})
                run = p.makeelement(qn('w:r'), {})
                if negrita:
                    rPr = run.makeelement(qn('w:rPr'), {})
                    rPr.append(rPr.makeelement(qn('w:b'), {}))
                    run.append(rPr)
                t = run.makeelement(qn('w:t'), {})
                t.set(qn('xml:space'), 'preserve')
                t.text = str(valor)
                run.append(t)
                p.append(run)
                tc.append(p)
                tr.append(tc)
            return tr

        nueva.append(fila(encabezados, negrita=True))
        for valores in filas:
            nueva.append(fila(valores))

        self.cursor.addnext(nueva)
        self.cursor = nueva
        # Word necesita un parrafo entre dos tablas seguidas; ademas separa visualmente.
        self.parrafo('')
        return nueva

    def imagen(self, ruta, pie=''):
        from PIL import Image
        with Image.open(ruta) as img:
            ancho_px, alto_px = img.size
        ancho = min(ANCHO_MAXIMO_CM, ANCHO_MAXIMO_CM)
        alto = ancho * alto_px / float(ancho_px)

        p = self.doc.add_paragraph()
        p.add_run().add_picture(ruta, width=Cm(ancho), height=Cm(alto))
        self.cursor.addnext(p._element)
        self.cursor = p._element
        if pie:
            self.parrafo(pie)
        return p._element

    def guardar(self, ruta):
        self.doc.save(ruta)
        podar(ruta)


def podar(ruta):
    """Saca del paquete las imagenes que ya no referencia ningun parrafo.

    Al reemplazar una seccion se borran sus parrafos, pero las imagenes que
    contenian siguen dentro del .docx y con su relacion declarada. Sin esta poda el
    archivo arrastra el peso de todas las figuras de la version anterior.
    """
    import re
    import shutil
    import zipfile

    with zipfile.ZipFile(ruta) as z:
        partes = {n: z.read(n) for n in z.namelist()}

    documento = partes['word/document.xml'].decode('utf-8')
    usadas = set(re.findall(r'r:(?:embed|id)="([^"]+)"', documento))

    rels = partes['word/_rels/document.xml.rels'].decode('utf-8')
    vivas, muertas = [], set()
    for rel in re.findall(r'<Relationship\b[^>]*/>', rels):
        ident = re.search(r'Id="([^"]+)"', rel).group(1)
        destino = re.search(r'Target="([^"]+)"', rel)
        es_imagen = destino and 'media/' in destino.group(1)
        if es_imagen and ident not in usadas:
            muertas.add('word/' + destino.group(1).lstrip('/'))
        else:
            vivas.append(rel)

    if not muertas:
        return

    nuevo_rels = re.sub(r'<Relationship\b[^>]*/>', '', rels)
    nuevo_rels = nuevo_rels.replace('</Relationships>',
                                    ''.join(vivas) + '</Relationships>')
    partes['word/_rels/document.xml.rels'] = nuevo_rels.encode('utf-8')

    # Una imagen puede estar referenciada por mas de una relacion viva.
    conservar = set()
    for rel in vivas:
        destino = re.search(r'Target="([^"]+)"', rel)
        if destino and 'media/' in destino.group(1):
            conservar.add('word/' + destino.group(1).lstrip('/'))
    muertas -= conservar

    temporal = ruta + '.tmp'
    with zipfile.ZipFile(temporal, 'w', zipfile.ZIP_DEFLATED) as z:
        for nombre, contenido in partes.items():
            if nombre in muertas:
                continue
            z.writestr(nombre, contenido)
    shutil.move(temporal, ruta)
    print('podadas %d imágenes sin referencia' % len(muertas))


# --------------------------------------------------------------------------- secciones

def seccion_diagramas_cu(d, casos, modulos):
    d.vaciar('2.2.1', '2.2.2')
    d.parrafo(
        'Se presenta un diagrama de casos de uso por cada módulo funcional del '
        'sistema. El actor principal es la encargada del sector, único usuario del '
        'sistema. El caso de uso de envío del resumen diario se dispara mediante un '
        'proceso programado, por lo que se representa con el actor Sistema, quedando '
        'la encargada como actor secundario destinatario del mensaje.')
    for numero in sorted(modulos):
        if not any(c['modulo'] == numero for c in casos):
            continue
        d.parrafo(modulos[numero], negrita=True)
        d.imagen(os.path.join(DIAGRAMAS, 'casos-de-uso', 'cu-modulo-%d.png' % numero),
                 'Figura. Diagrama de casos de uso — %s.' % modulos[numero])


def seccion_casos_de_uso(d, casos, modulos):
    d.vaciar('2.2.2', '2.2.3')
    d.parrafo(
        'Un caso de uso es una herramienta que sirve para representar la forma en que '
        'un actor interactúa con el sistema para alcanzar un objetivo determinado. A '
        'continuación se listan los cuarenta y nueve casos de uso identificados, '
        'agrupados por módulo, y luego se detalla cada uno de ellos.')

    modulo_actual = None
    for caso in casos:
        if caso['modulo'] != modulo_actual:
            modulo_actual = caso['modulo']
            d.parrafo(modulos[modulo_actual], negrita=True)
        d.parrafo('CU %d — %s' % (caso['num'], caso['nombre']))

    modulo_actual = None
    for caso in casos:
        if caso['modulo'] != modulo_actual:
            modulo_actual = caso['modulo']
            d.parrafo(modulos[modulo_actual], negrita=True)
        for i, linea in enumerate(lineas(caso)):
            d.parrafo(linea, negrita=(i == 0))


def seccion_dominio(d):
    d.vaciar('2.2.3', '2.2.4')
    d.parrafo(
        'El diagrama de dominio presenta las clases de negocio del sistema y la clase '
        'Controladora, que concentra la lógica. Las clases de negocio son anémicas: '
        'conservan sus atributos y su constructor, y no contienen reglas. Hembra y '
        'Macho especializan a Animal y comparten su identidad. Cuatro clases '
        '—ControlDiario, CandidataDescarte, PartidaVencimiento y '
        'ProcedimientoPendiente— no tienen tabla en la base: transportan un resultado '
        'que el sistema deriva en el momento de la consulta.')
    d.imagen(os.path.join(DIAGRAMAS, 'dominio-clases-de-negocio.png'),
             'Figura. Diagrama de Dominio — clases de negocio.')
    d.parrafo(
        'La Controladora agrupa sus operaciones por región funcional. Valida contra '
        'sus listas en memoria, delega la escritura en pControladora y recién después '
        'actualiza la caché. Es además el único punto de contacto entre la capa de '
        'dominio y la capa de acceso a datos.')
    d.imagen(os.path.join(DIAGRAMAS, 'dominio-controladora.png'),
             'Figura. Diagrama de Dominio — clase Controladora.')


def seccion_persistencia(d):
    d.vaciar('2.2.4', '2.2.5')
    d.parrafo(
        'La capa de acceso a datos se organiza detrás de la fachada pControladora, '
        'que es la única clase que la Controladora conoce. Cada operación deriva en la '
        'clase de persistencia de la entidad correspondiente, y todas resuelven su '
        'ejecución contra la base a través de pConexion.')
    d.imagen(os.path.join(DIAGRAMAS, 'persistencia-fachada.png'),
             'Figura. Diagrama de Persistencia — fachada pControladora.')
    d.imagen(os.path.join(DIAGRAMAS, 'persistencia-clases.png'),
             'Figura. Diagrama de Persistencia — clases de acceso a datos.')


def seccion_modelo_datos(d):
    d.vaciar('2.2.5.1', '2.2.5.2')
    d.parrafo(
        'El modelo entidad-relación reúne las veinticuatro relaciones del esquema. Las '
        'dos últimas —preferencias_notificacion y alertas— corresponden al Módulo 7, y '
        'no guardan información que el sistema no tenga: registran a qué avisos adhirió '
        'el establecimiento y cuáles se enviaron, porque los pendientes que se avisan se '
        'derivan de las relaciones anteriores.')
    d.imagen(os.path.join(DIAGRAMAS, 'modelo-entidad-relacion.png'),
             'Figura. Modelo Entidad-Relación.')

    d.vaciar('2.2.5.2', '2.2.5.3')
    d.parrafo(
        'A partir del Modelo Entidad-Relación se derivó el esquema relacional que se '
        'presenta a continuación. Cada entidad da origen a una relación y cada vínculo '
        'uno a muchos se resuelve mediante la clave foránea correspondiente del lado de '
        'la cardinalidad muchos. Los dos vínculos muchos a muchos del esquema originan '
        'relaciones intermedias: plan_categorias, entre planes_sanitarios y categorias, '
        'y ordenie_lote_animales, entre ordenies_lote y hembras. Las jerarquías de '
        'especialización de Animal se resuelven por tablas separadas que comparten la '
        'clave primaria de la entidad padre.')
    d.parrafo(
        'El esquema resultante se encuentra en Tercera Forma Normal. Se verifica la '
        'Primera Forma Normal dado que todos los atributos son atómicos y no existen '
        'grupos repetitivos. Se verifica la Segunda Forma Normal porque las dos '
        'relaciones con clave primaria compuesta —plan_categorias y '
        'ordenie_lote_animales— no poseen atributos no clave y por lo tanto no admiten '
        'dependencias parciales; en las relaciones restantes la clave primaria está '
        'formada por un único atributo. Se verifica la Tercera Forma Normal porque '
        'ningún atributo no clave depende de otro atributo no clave.')
    d.parrafo(
        'Merece una aclaración el tratamiento de los valores que el sistema calcula. La '
        'categoría del animal, la fecha probable de parto y la fecha de fin de descarte '
        'de leche se proponen de forma automática a partir de datos ya registrados, '
        'pero en los tres casos el usuario puede corregir el valor propuesto. Esa '
        'posibilidad de ajuste rompe la dependencia funcional con los atributos que les '
        'dieron origen: el dato almacenado deja de ser derivable y pasa a ser un dato '
        'propio de la relación, por lo que su persistencia no constituye una '
        'redundancia. En cambio, los valores que el sistema calcula y el usuario no '
        'puede modificar —los procedimientos pendientes del calendario sanitario, el '
        'remanente de cada partida, los indicadores del rodeo o la producción estimada '
        'de una lactancia— no se almacenan y se derivan en el momento de la consulta.')
    d.parrafo('Se subrayan las claves primarias de cada relación.')
    for nombre, columnas in modelo_datos.normalizacion():
        d.parrafo('%s = {%s}' % (nombre, ', '.join(columnas)))

    d.vaciar('2.2.5.3', '2.2.5.4')
    d.parrafo(
        'Se detallan, para cada tabla del esquema, su clave primaria, sus claves '
        'alternas y sus claves foráneas.')
    d.parrafo(
        'En las tablas hembras y machos la clave primaria id_animal es a su vez clave '
        'foránea hacia animales, dado que ambas resuelven la especialización de la '
        'entidad Animal. En plan_categorias y en ordenie_lote_animales la clave '
        'primaria es compuesta y está formada por las dos claves foráneas que resuelven '
        'el vínculo muchos a muchos.')
    d.tabla(['Tabla', 'Claves primarias', 'Claves alternas', 'Claves foráneas'],
            modelo_datos.claves())

    d.vaciar('2.2.5.4', '2.2.6')
    d.parrafo(
        'Para cada tabla del esquema se detallan sus campos con el tipo de dato '
        'definido, las restricciones que se aplican sobre ellos y las observaciones '
        'pertinentes. Los identificadores se implementan como enteros auto '
        'incrementados, con la excepción de plan_categorias y ordenie_lote_animales, '
        'cuyas claves primarias están formadas por sus dos claves foráneas. Las claves '
        'foráneas que admiten valor nulo corresponden a vínculos opcionales, como los '
        'progenitores no registrados, el tratamiento preventivo sin diagnóstico o la '
        'aplicación registrada fuera de plan.')
    for nombre, filas in modelo_datos.integridad():
        d.parrafo('Tabla: %s' % nombre, negrita=True)
        d.tabla(['Campo', 'Tipo', 'Restricciones', 'Observaciones'], filas)


def seccion_secuencia(d, casos):
    d.vaciar('2.2.6', '2.2.7')
    d.parrafo(
        'Se presenta un diagrama de secuencia por cada uno de los cuarenta y nueve '
        'casos de uso. Todos siguen el mismo recorrido entre capas: el actor interactúa '
        'con la vista del módulo, la vista invoca a la Controladora, y ésta valida '
        'contra sus listas en memoria antes de delegar la escritura en pControladora, '
        'que a su vez deriva en la clase de persistencia de la entidad y en pConexion. '
        'Las validaciones que se resuelven en memoria aparecen como auto-mensajes sobre '
        'la Controladora.')
    d.parrafo(
        'Tres casos se apartan del patrón. El CU 1 no involucra la capa de persistencia, '
        'porque el sistema opera con un único par de credenciales fijas. Los CU 48 y '
        'CU 49 incorporan la línea de vida BotTelegram, y el mensaje sale desde la vista '
        'y desde el proceso, no desde la Controladora: ésta arma el texto, pero salir a '
        'internet no es parte del dominio. El CU 49, además, no tiene vista, porque su '
        'actor es el proceso programado que dispara el envío del resumen diario.')
    for caso in casos:
        d.parrafo('CU %d — %s' % (caso['num'], caso['nombre']), negrita=True)
        d.imagen(os.path.join(DIAGRAMAS, 'secuencia', 'secuencia-cu%02d.png' % caso['num']),
                 'Figura. Diagrama de secuencia del CU %d.' % caso['num'])


def seccion_diccionario(d):
    d.vaciar('2.2.7', '2.3')
    d.parrafo(
        'Se especifican los tipos que toman los atributos de cada clase de la capa de '
        'dominio, junto con los parámetros que reciben sus métodos y una descripción de '
        'su cometido. Las clases Hembra y Macho heredan de Animal, por lo que sólo se '
        'detallan los atributos propios de cada especialización. Cada atributo privado '
        'mX se expone mediante una propiedad pública X con getter y setter. Las clases '
        'de dominio no contienen lógica de negocio: son anémicas y conservan únicamente '
        'sus atributos y su constructor. Toda la lógica reside en la clase Controladora, '
        'que valida contra sus listas static en memoria, delega la escritura en '
        'pControladora y recién después actualiza la caché.')

    d.parrafo('Clases de negocio', negrita=True)
    for nombre, base, nota, filas in diccionario_clases.entidades():
        titulo = nombre + (' (hereda de %s)' % base if base else '')
        d.parrafo(titulo, negrita=True)
        if nota:
            d.parrafo(nota)
        d.tabla(['Atributo', 'Tipo', 'Descripción'], filas)

    d.parrafo('Controladora', negrita=True)
    d.parrafo(
        'Los métodos se agrupan por región funcional, en el mismo orden en que están '
        'escritos en la clase.')
    for region, metodos in diccionario_clases.controladora():
        d.parrafo('%s (%d métodos)' % (region.title(), len(metodos)), negrita=True)
        d.tabla(['Método', 'Devuelve', 'Parámetros'],
                [(n, t, p or '—') for n, t, p in metodos])

    d.parrafo('pControladora', negrita=True)
    d.parrafo(
        'La fachada de la capa de acceso a datos. Cada método deriva en la clase de '
        'persistencia de su entidad.')
    for region, metodos in diccionario_clases.fachada():
        d.parrafo('%s (%d métodos)' % (region.title(), len(metodos)), negrita=True)
        d.tabla(['Método', 'Devuelve', 'Parámetros'],
                [(n, t, p or '—') for n, t, p in metodos])

    d.parrafo('Clases de acceso a datos', negrita=True)
    for nombre, metodos in diccionario_clases.acceso_a_datos():
        d.parrafo(nombre, negrita=True)
        d.tabla(['Método', 'Devuelve', 'Parámetros'],
                [(n, t, p or '—') for n, t, p in metodos])


def seccion_analisis(d):
    """Actualiza los dos parrafos de 2.1 que describen el alcance del sistema."""
    i = d.buscar('El sistema será una aplicación web operada por un único usuario')
    elemento = list(d.cuerpo.iterchildren())[i]
    d.cursor = elemento
    nuevo = d.parrafo(
        'El sistema es una aplicación web operada por un único usuario con credenciales '
        'fijas, por lo que no requiere gestión de roles ni administración de múltiples '
        'cuentas. Su funcionalidad se organiza en ocho módulos: acceso al sistema y '
        'configuración de los parámetros de manejo del establecimiento; gestión de '
        'animales y genética (altas y bajas, linaje, control de consanguinidad y '
        'clasificación por categoría); producción y ordeñe (registro por lote, control '
        'lechero y seguimiento de lactancias); reproducción (celos, servicios, tactos y '
        'partos); sanidad (diagnósticos, tratamientos, vacunaciones y descornes); '
        'insumos y stock (movimientos, control de stock mínimo y vencimientos por '
        'partida); tablero, indicadores y apoyo a la decisión de descarte; y, '
        'finalmente, reportes y notificaciones.')
    d.cuerpo.remove(elemento)
    return nuevo


# --------------------------------------------------------------------------- principal

def main():
    casos, modulos = cargar()
    d = Documento(ENTRADA)

    seccion_analisis(d)
    seccion_diagramas_cu(d, casos, modulos)
    seccion_casos_de_uso(d, casos, modulos)
    seccion_dominio(d)
    seccion_persistencia(d)
    seccion_modelo_datos(d)
    seccion_secuencia(d, casos)
    seccion_diccionario(d)

    d.guardar(SALIDA)
    print('guardado', SALIDA)


if __name__ == '__main__':
    main()
