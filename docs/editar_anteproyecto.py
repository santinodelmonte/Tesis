# -*- coding: utf-8 -*-
"""Produce Anteproyecto_v7.docx aplicando sobre el v5 dos tandas de cambios.

El v6 dice lo que el sistema hace: cada requerimiento que este script reescribe o
agrega esta verificado contra Tesis/Dominio, Tesis/Persistencia, Tesis/Pages y
bd/CreacionDb.sql. El resumen de la primera tanda esta en docs/cambios-anteproyecto-v6.md.

La segunda tanda -el bloque v7 del final- corrige lo que encontro la revision de
ingenieria de software de docs/revision-tutor.md: objetivos que prometian porcentajes
sin linea base, requerimientos que no se podian verificar y actores del sistema
mezclados con interesados del proyecto. El resumen esta en
docs/cambios-anteproyecto-v7.md.

Trabaja sobre el .docx original conservando estilos, numeraciones y formato: los
requerimientos nuevos se crean clonando un requerimiento existente del mismo modulo,
asi heredan la lista numerada y la sangria que ya tenia el documento.
"""
import copy
from docx import Document
from docx.oxml.ns import qn

import os
AQUI = os.path.dirname(os.path.abspath(__file__))
RAIZ = os.path.dirname(AQUI)
RUTA_ENTRADA = os.path.join(RAIZ, 'Anteproyecto_v5.docx')
RUTA_SALIDA = os.path.join(RAIZ, 'Anteproyecto_v7.docx')

doc = Document(RUTA_ENTRADA)
BULLET = '●​ '


# --------------------------------------------------------------------------- utilidades

def buscar(prefijo, desde=0):
    """Devuelve el parrafo cuyo texto arranca con el prefijo dado."""
    for i, p in enumerate(doc.paragraphs):
        if i >= desde and p.text.strip().startswith(prefijo):
            return p
    raise LookupError('no se encontro el parrafo: ' + prefijo)


def primer_run(p):
    for r in p.runs:
        return r
    raise LookupError('parrafo sin runs: ' + p.text[:40])


def escribir(p, titulo, cuerpo=None, salto=True):
    """Reescribe el contenido de un parrafo conservando su formato.

    Con salto=True reproduce la forma "titulo:<br> cuerpo" que usan los modulos 0, 1,
    2, 5 y 6; con salto=False la forma "titulo: cuerpo" en una sola linea de los
    modulos 3 y 4.
    """
    run = primer_run(p)
    for otro in p.runs[1:]:
        otro._element.getparent().remove(otro._element)

    elem = run._element
    for hijo in list(elem):
        if hijo.tag != qn('w:rPr'):
            elem.remove(hijo)

    def texto(valor):
        t = elem.makeelement(qn('w:t'), {})
        t.set(qn('xml:space'), 'preserve')
        t.text = valor
        elem.append(t)

    if cuerpo is None:
        texto(titulo)
    elif salto:
        texto(titulo)
        br = elem.makeelement(qn('w:br'), {})
        br.set(qn('w:type'), 'textWrapping')
        elem.append(br)
        texto(' ' + cuerpo)
    else:
        texto(titulo + ' ' + cuerpo)


def clonar_despues(modelo, titulo, cuerpo=None, salto=True):
    """Inserta una copia del parrafo modelo justo despues de el, con texto nuevo."""
    return _clonar(modelo, modelo, 'next', titulo, cuerpo, salto)


def clonar_antes(modelo, referencia, titulo, cuerpo=None, salto=True):
    """Copia el formato del modelo y coloca el parrafo nuevo antes de la referencia."""
    return _clonar(modelo, referencia, 'prev', titulo, cuerpo, salto)


def _clonar(modelo, referencia, donde, titulo, cuerpo, salto):
    nuevo = copy.deepcopy(modelo._element)
    if donde == 'next':
        referencia._element.addnext(nuevo)
    else:
        referencia._element.addprevious(nuevo)
    p = [x for x in doc.paragraphs if x._element is nuevo][0]
    escribir(p, titulo, cuerpo, salto)
    return p


def agregar_a(p, frase):
    """Suma una frase al final del ultimo run del parrafo."""
    run = p.runs[-1]
    ts = run._element.findall(qn('w:t'))
    ultimo = ts[-1]
    ultimo.set(qn('xml:space'), 'preserve')
    ultimo.text = (ultimo.text or '') + frase


# --------------------------------------------------------- Modulo 0: acceso y configuracion

escribir(buscar('Módulo 0:'), 'Módulo 0: Seguridad, Acceso y Configuración:')

p = buscar('RF0.1')
p = clonar_despues(
    p, 'RF0.2 Cierre de sesión:',
    'El sistema debe permitir cerrar la sesión activa desde cualquier pantalla, '
    'impidiendo el acceso posterior a la información sin volver a autenticarse.')
clonar_despues(
    p, 'RF0.3 Configuración de parámetros de manejo:',
    'El sistema debe permitir configurar los parámetros de manejo del establecimiento '
    '—días de secado previos al parto, edad mínima al servicio, edad de cambio de '
    'categoría, litros máximos admitidos por control, cantidad de ordeñes diarios, '
    'espera voluntaria posparto, días para el tacto y la anticipación de cada uno de '
    'los cuatro avisos—, validando que cada valor quede dentro de su rango admitido y '
    'aplicando valores por defecto mientras no se configuren.')


# ------------------------------------------------------ Modulo 1: animales y genetica

escribir(
    buscar('RF1.2'), 'RF1.2 Baja de animales:',
    'El sistema debe permitir registrar la baja lógica de un animal indicando el motivo '
    'de salida (venta, fallecimiento, descarte sanitario u otros), dejando asentada la '
    'fecha en que la baja se registra y conservando su historial y su lugar en el linaje '
    'del rodeo.')

# La verificacion no se limita al parentesco directo: recorre la ascendencia registrada
# de los dos animales y avisa si encuentran un ancestro comun.
escribir(
    buscar('RF1.7'), 'RF1.7 Prevención de consanguinidad:',
    'El sistema debe advertir el parentesco entre una hembra y un posible reproductor, '
    'verificando si existe un ancestro común dentro de la ascendencia registrada de '
    'ambos. La advertencia es informativa y no impide registrar el servicio.')

# El macho de mas de quince meses es Toro si se destina a la reproduccion -esta en pie o
# aporta pajuelas al stock-; si no, es Novillo. La edad sola no alcanza para
# clasificarlo, y la presencia en el campo tampoco: el toro de catalogo no integra el
# rodeo y sigue siendo un reproductor.
escribir(
    buscar('RF1.8'), 'RF1.8 Clasificación automática:',
    'El sistema debe clasificar automáticamente a los animales: las hembras según su '
    'edad y su número de partos, y los machos según su edad y su destino reproductivo.')

escribir(
    buscar('RF1.5'), 'RF1.5 Registro genealógico:',
    'El sistema debe permitir registrar el padre y la madre de cada animal al momento '
    'del alta, admitiendo que alguno de ellos no se encuentre registrado en el sistema.')

escribir(
    buscar('RF1.9'), 'RF1.9 Actualización de categoría:',
    'El sistema debe recalcular la categoría que corresponde a cada animal cuando cambia '
    'su condición biológica, señalando en el listado y en la ficha aquellos cuya '
    'categoría quedó desactualizada y permitiendo aceptar la nueva. La actualización se '
    'aplica de forma automática al registrarse un parto.')

escribir(
    buscar('RF1.10'), 'RF1.10 Consulta y búsqueda:',
    'El sistema debe permitir búsquedas y filtros combinados por número de caravana, '
    'raza, categoría, estado y rango etario.')

p = buscar('RF1.10')
p = clonar_despues(
    p, 'RF1.11 Reactivación de animales:',
    'El sistema debe permitir revertir la baja de un animal, devolviéndolo al rodeo y '
    'limpiando su fecha y su motivo de salida.')
p = clonar_despues(
    p, 'RF1.12 Fotografía del animal:',
    'El sistema debe permitir asociar una fotografía a cada animal y mostrarla en su '
    'ficha y en el árbol genealógico.')
p = clonar_despues(
    p, 'RF1.13 Árbol genealógico interactivo:',
    'El sistema debe presentar la ascendencia de cada animal como un árbol navegable, '
    'permitiendo desplegar cada rama y acceder a la ficha de cualquier ancestro.')
p = clonar_despues(
    p, 'RF1.14 Validación del árbol genealógico:',
    'El sistema debe impedir el registro de una genealogía imposible —un animal como '
    'progenitor de sí mismo, un progenitor que figure en su propia descendencia, o un '
    'progenitor cuya fecha de nacimiento no admita la edad mínima al servicio— y debe '
    'advertir cuando el progenitor elegido se encuentre dado de baja.')
clonar_despues(
    p, 'RF1.15 Ficha integral del animal:',
    'El sistema debe presentar en una sola pantalla la identificación, la categoría, el '
    'linaje y el historial productivo, reproductivo y sanitario de cada animal.')


# ------------------------------------------------------- Modulo 2: control de produccion
# RF2.1 y RF2.2 se unifican en un unico requerimiento y el resto del modulo se corre un
# lugar: los dos ordenies diarios dejaron de ser una constante del sistema.

escribir(
    buscar('RF2.1'), 'RF2.1 Registro de ordeñe por lote:',
    'El sistema debe permitir registrar los litros totales producidos por el rodeo en '
    'cada turno de ordeñe de la jornada, según la cantidad de ordeñes diarios '
    'configurada.')

viejo = buscar('RF2.2')
viejo._element.getparent().remove(viejo._element)

escribir(
    buscar('RF2.3'), 'RF2.2 Validación de producción:',
    'El sistema debe validar que los litros ingresados sean valores positivos, que no '
    'superen el máximo configurado por control individual, y que en el registro por lote '
    'no superen ese máximo multiplicado por la cantidad de animales del lote.')

escribir(
    buscar('RF2.4'), 'RF2.3 Control lechero:',
    'El sistema debe permitir registrar la producción individual de cada animal en un '
    'turno determinado, admitiendo tanto la carga masiva de todo el rodeo en ordeñe como '
    'la carga puntual de un solo animal. El control lechero mide al animal y convive con '
    'el registro por lote, sin sumarse a él dentro de un mismo turno.')

escribir(
    buscar('RF2.5'), 'RF2.4 Total diario general:',
    'El sistema debe calcular la producción total diaria del establecimiento.')

escribir(
    buscar('RF2.6'), 'RF2.5 Historial productivo individual:',
    'El sistema debe almacenar el histórico de los controles individuales registrados '
    'para cada animal.')

escribir(
    buscar('RF2.7'), 'RF2.6 Historial de lactancias:',
    'El sistema debe registrar y visualizar la producción histórica organizada por '
    'lactancia, estimando el total producido a partir de los intervalos entre controles '
    'y proyectándolo a 305 días. La proyección es lineal: permite comparar animales '
    'entre sí, sin constituir un pronóstico.')

escribir(
    buscar('RF2.8'), 'RF2.7 Producción entre fechas:',
    'El sistema debe calcular la producción total e individual en un rango de fechas '
    'seleccionado.')

escribir(
    buscar('RF2.9'), 'RF2.8 Registro de secado:',
    'El sistema debe permitir registrar manualmente el inicio del período de secado.')

escribir(
    buscar('RF2.10'), 'RF2.9 Cálculo automático de secado:',
    'El sistema debe calcular la fecha recomendada de secado a partir de la fecha '
    'probable de parto y de los días de secado configurados.')

escribir(
    buscar('RF2.11'), 'RF2.10 Alertas de secado:',
    'El sistema debe generar alertas para los animales próximos al secado, con la '
    'anticipación configurada.')

escribir(
    buscar('RF2.12'), 'RF2.11 Cambio de estado productivo:',
    'El sistema debe actualizar automáticamente el estado productivo del animal: a '
    '“seca” cuando se registra el secado, y a “en lactancia” cuando se registra el parto.')

p = buscar('RF2.11 Cambio')
p = clonar_despues(
    p, 'RF2.12 Apertura manual de lactancia:',
    'El sistema debe permitir abrir una lactancia sin un parto registrado, para los '
    'animales que ya se encontraban en ordeñe al comenzar a utilizarse el sistema.')
clonar_despues(
    p, 'RF2.13 Corrección y eliminación de registros de producción:',
    'El sistema debe permitir corregir y eliminar los ordeñes por lote y los controles '
    'lecheros ya registrados, informando qué registros dependen de aquel que se pretende '
    'eliminar.')


# --------------------------------------------------------- Modulo 3: gestion reproductiva

escribir(
    buscar('RF3.8'), 'RF3.8 Registro de parto:',
    'El sistema debe registrar la fecha del parto, su tipo y las observaciones, y dar de '
    'alta cada cría como animal del rodeo con su caravana, raza y fotografía, proponiendo '
    'como padre el toro del servicio que originó la preñez.',
    salto=False)

p = buscar('RF3.9')
p = clonar_despues(
    p, 'RF3.10 Validaciones reproductivas:',
    'El sistema debe impedir el registro de un celo o de un servicio en un animal que no '
    'alcanzó la edad mínima al servicio, y el de cualquier evento con fecha posterior a '
    'la baja del animal.', salto=False)
p = clonar_despues(
    p, 'RF3.11 Advertencias reproductivas:',
    'El sistema debe advertir, sin impedir el registro, ante situaciones inusuales pero '
    'posibles: servicio con un toro dado de baja, parentesco entre la hembra y el '
    'reproductor, mellizos de distinto sexo, duración de la gestación fuera del rango '
    'normal y parto de una hembra que no figuraba preñada.', salto=False)
p = clonar_despues(
    p, 'RF3.12 Listas de trabajo:',
    'El sistema debe presentar la lista de los servicios con tacto pendiente y la de las '
    'hembras en condiciones de ser servidas, indicando en cada caso el motivo por el cual '
    'figuran.', salto=False)
clonar_despues(
    p, 'RF3.13 Corrección y eliminación de eventos reproductivos:',
    'El sistema debe permitir corregir y eliminar celos, servicios, tactos y partos, '
    'volviendo a deducir el estado reproductivo del animal a partir de los eventos que '
    'permanecen registrados.', salto=False)


# ------------------------------------------------------------ Modulo 4: gestion sanitaria

p = buscar('RF4.7')
p = clonar_despues(
    p, 'RF4.8 Cierre del diagnóstico:',
    'El sistema debe permitir dar por resuelto un diagnóstico, distinguiendo los cuadros '
    'sanitarios activos de los ya cerrados.', salto=False)
clonar_despues(
    p, 'RF4.9 Corrección y eliminación de eventos sanitarios:',
    'El sistema debe permitir corregir y eliminar diagnósticos, tratamientos, '
    'vacunaciones y descornes, devolviendo al stock los insumos que se habían '
    'descontado.', salto=False)


# ------------------------------------------------------ Modulo 5: control de insumos y stock

escribir(
    buscar('RF5.3'), 'RF5.3 Descuento de medicamentos:',
    'El sistema debe descontar del stock la cantidad de insumo aplicada al registrarse un '
    'tratamiento. La cantidad se ingresa al registrar el tratamiento, ya que depende de '
    'la dosis indicada y de la presentación del producto.')

escribir(
    buscar('RF5.7'), 'RF5.7 Fecha de vencimiento:',
    'El sistema debe almacenar el vencimiento de cada partida ingresada de medicamentos y '
    'vacunas, llevando el stock del insumo desagregado por partida.')

escribir(
    buscar('RF5.8'), 'RF5.8 Alertas de vencimiento:',
    'El sistema debe notificar los próximos vencimientos con la anticipación configurada, '
    'imputando el consumo a la partida que vence primero.')

clonar_despues(
    buscar('RF5.9'), 'RF5.10 Reversión de movimientos de stock:',
    'Al corregirse o eliminarse el evento que consumió un insumo, el sistema debe '
    'devolver la cantidad mediante un contra-movimiento, conservando el movimiento '
    'original en el historial.')


# --------------------------- Modulo 6 nuevo: tablero e indicadores; reportes pasa al 7

# Primero el modulo de reportes pasa a ser el septimo, con sus requerimientos
# renumerados. Recien despues se crea el sexto: si se hiciera al reves, los RF6.x
# nuevos y los viejos convivirian con el mismo nombre.
escribir(buscar('Módulo 6: Reportes'), 'Módulo 7: Reportes y Notificaciones:')

reportes = [
    ('RF6.1', 'RF7.1 Reportes productivos:',
     'El sistema debe generar reportes PDF y Excel de producción individual y general.'),
    ('RF6.2', 'RF7.2 Reportes sanitarios:',
     'El sistema debe generar reportes de enfermedades, tratamientos y vacunaciones.'),
    ('RF6.3', 'RF7.3 Reportes reproductivos:',
     'El sistema debe generar reportes sobre servicios, preñeces, partos y secados.'),
    ('RF6.4', 'RF7.4 Reportes genéticos:',
     'El sistema debe generar reportes de genealogía y rendimiento por línea genética.'),
    ('RF6.5', 'RF7.5 Integración con Telegram:',
     'El sistema debe integrarse con un bot de Telegram.'),
    ('RF6.6', 'RF7.6 Notificaciones automáticas:',
     'El sistema debe enviar alertas automáticas sobre procedimientos sanitarios '
     'pendientes, partos próximos, tactos pendientes, secados próximos, stock crítico, '
     'vencimiento de insumos y fin del período de descarte de leche.'),
    ('RF6.7', 'RF7.7 Resumen diario:',
     'El sistema debe enviar un resumen diario de las tareas pendientes.'),
]
for prefijo, titulo, cuerpo in reportes:
    escribir(buscar(prefijo), titulo, cuerpo)

# El modulo nuevo se arma entre el modulo 5 y el de reportes: cada requerimiento se
# inserta delante del encabezado del septimo, de modo que quedan en orden.
encabezado_reportes = buscar('Módulo 7: Reportes')
modelo = buscar('RF7.1')
encabezado_tablero = copy.deepcopy(encabezado_reportes._element)
encabezado_reportes._element.addprevious(encabezado_tablero)
p = [x for x in doc.paragraphs if x._element is encabezado_tablero][0]
escribir(p, 'Módulo 6: Tablero, Indicadores y Apoyo a la Decisión:')

p = clonar_antes(
    modelo, encabezado_reportes, 'RF6.1 Tablero de inicio:',
    'El sistema debe presentar, como pantalla de entrada, el estado del día del '
    'establecimiento: las tareas pendientes y los avisos vigentes de cada módulo.')
p = clonar_despues(
    p, 'RF6.2 Indicadores del rodeo:',
    'El sistema debe calcular y presentar los indicadores de desempeño del rodeo: días '
    'abiertos, intervalo entre partos, servicios por preñez, litros por vaca y por día, '
    'días en leche promedio, composición del rodeo por estado productivo y reproductivo, '
    'y el ranking de las lactancias en curso con su proyección a 305 días.')
p = clonar_despues(
    p, 'RF6.3 Candidatas a descarte:',
    'El sistema debe listar las hembras que cumplen al menos uno de los criterios de '
    'descarte —producción por debajo del 70 % del promedio del rodeo, tres o más '
    'servicios sin preñez, más de 150 días abiertos, tres o más diagnósticos en el último '
    'año, o siete o más partos—, indicando en cada caso los motivos por los que figuran. '
    'El sistema informa: la decisión es de la encargada.')
clonar_despues(
    p, 'RF6.4 Buscador de caravana:',
    'El sistema debe ofrecer, desde cualquier pantalla, un buscador que lleve '
    'directamente a la ficha del animal a partir de su número de caravana.')


# -------------------------------------------------------- requerimientos no funcionales

agregar_a(
    buscar('El sistema debe presentar una interfaz simple'),
    ' La navegación se organiza en un menú por módulo que presenta primero el listado de '
    'cada entidad, los listados extensos se paginan y un buscador de caravana permanente '
    'permite llegar a cualquier animal desde cualquier pantalla.')

agregar_a(
    buscar('El sistema debe ser accesible a través de navegadores'),
    ' Asimismo debe cumplir con las pautas WCAG 2.1 en su nivel AA, verificando los '
    'contrastes de la paleta y evitando que el color sea el único medio para transmitir '
    'información.')

agregar_a(
    buscar('El sistema debe ser responsive'),
    ' El uso desde el celular debe estar verificado a partir de los 375 píxeles de ancho, '
    'con las tablas adaptadas a la lectura en pantalla angosta.')

agregar_a(
    buscar('El sistema debe garantizar la integridad de los datos'),
    ' Las operaciones que escriben sobre más de una tabla deben resolverse dentro de una '
    'transacción, de modo que un fallo parcial no deje registros incompletos.')

agregar_a(
    buscar('El sistema debe restringir el acceso a la información'),
    ' Las credenciales de acceso y la cadena de conexión a la base de datos no deben '
    'residir en el código fuente, y las consultas deben construirse de forma '
    'parametrizada.')


# ------------------------------------------------------------------ alcance y limitaciones

p = buscar(BULLET + 'Apoyo a la toma de decisiones')
p = clonar_despues(
    p, BULLET + 'Configuración de los parámetros de manejo propios del establecimiento.')
p = clonar_despues(
    p, BULLET + 'Tablero de tareas del día, indicadores del rodeo y apoyo a la decisión '
    'de descarte.')
clonar_despues(
    p, BULLET + 'Corrección y eliminación de los registros ya cargados, con reversión de '
    'sus efectos.')

p = buscar(BULLET + 'No se contempla la migración completa')
p = clonar_despues(
    p, BULLET + 'La configuración de parámetros es única para todo el establecimiento: no '
    'admite un manejo diferenciado por grupo de animales.')
p = clonar_despues(
    p, BULLET + 'El paginado de los listados se resuelve en el navegador, lo que resulta '
    'adecuado para el volumen de un establecimiento de esta escala.')
p = clonar_despues(
    p, BULLET + 'La proyección de producción a 305 días es lineal y no contempla la curva '
    'de lactancia.')
clonar_despues(
    p, BULLET + 'Los umbrales que determinan las candidatas a descarte son criterios fijos '
    'del sistema y no parámetros configurables.')


# ---------------------------------------------------------------------------- iteraciones

escribir(
    buscar('El desarrollo será dividido en seis incrementos'),
    'El desarrollo fue dividido en seis incrementos principales. Los cinco primeros se '
    'encuentran completados y el sexto es el que resta.')

iteraciones = [
    ('Primera Iteración',
     'En la primera iteración se desarrollaron el acceso al sistema y la gestión del '
     'rodeo, que es la base sobre la que se apoya todo el resto.',
     ['Diseño inicial de interfaces.',
      'Login con credenciales únicas.',
      'Alta, baja, modificación y búsqueda de animales.',
      'Registro genealógico, consulta de linaje y control de consanguinidad.',
      'Clasificación automática por categoría.',
      'Pruebas funcionales.']),
    ('Segunda Iteración',
     'En la segunda iteración se desarrollaron el control productivo y el manejo '
     'reproductivo. Se abordaron juntos porque el estado productivo de una hembra '
     'depende de sus partos y de sus secados: separarlos obligaba a construir dos veces '
     'la misma lógica de estados.',
     ['Ordeñe por lote y control lechero individual.',
      'Lactancias, secado y alertas de secado.',
      'Celos, servicios, tactos y partos.',
      'Cálculo de la fecha probable de parto y alertas.',
      'Pruebas funcionales.']),
    ('Tercera Iteración',
     'En la tercera iteración se incorporaron la gestión sanitaria y el control de '
     'insumos, que dependen entre sí: todo tratamiento consume stock.',
     ['Diagnósticos, tratamientos, vacunaciones y descornes.',
      'Planes sanitarios configurables y calendario sanitario.',
      'Alta de insumos, ingresos y movimientos de stock.',
      'Stock mínimo, vencimientos y período de descarte de leche.',
      'Pruebas funcionales.']),
    ('Cuarta Iteración',
     'En la cuarta iteración se incorporaron las reglas de negocio transversales que el '
     'uso real del sistema hizo evidentes.',
     ['Validaciones bloqueantes y advertencias no bloqueantes.',
      'Parámetros de manejo configurables por establecimiento.',
      'Baja reversible de animales.',
      'Paginado de los listados.',
      'Pruebas funcionales.']),
    ('Quinta Iteración',
     'En la quinta iteración se trabajó sobre lo que convierte a los datos registrados en '
     'información útil para la jornada, y sobre las condiciones de uso reales del '
     'establecimiento.',
     ['Tablero de inicio, indicadores del rodeo y listas de trabajo.',
      'Candidatas a descarte y buscador de caravana.',
      'Corrección y eliminación de registros ya cargados.',
      'Estilos, accesibilidad y uso desde el celular.',
      'Reorganización de la navegación.',
      'Pruebas funcionales.']),
    ('Sexta Iteración',
     'En la sexta y última iteración se desarrollarán los reportes y las notificaciones, '
     'junto con las validaciones finales del sistema.',
     ['Reportes productivos, sanitarios, reproductivos y genéticos.',
      'Integración con el bot de Telegram.',
      'Notificaciones automáticas y resumen diario de tareas.',
      'Validaciones integrales del sistema.',
      'Testing general y corrección de errores finales.']),
]

for encabezado, introduccion, items in iteraciones:
    parrafos = doc.paragraphs
    inicio = next(i for i, x in enumerate(parrafos)
                  if x.text.strip().startswith(encabezado))
    escribir(parrafos[inicio + 1], introduccion)
    escribir(parrafos[inicio + 2], 'Las funcionalidades comprendidas son:')

    viejos = []
    for siguiente in parrafos[inicio + 3:]:
        if siguiente.style.name.startswith('Heading') or not siguiente.text.strip():
            break
        viejos.append(siguiente)

    modelo = viejos[0]
    for texto_item, destino in zip(items, viejos):
        escribir(destino, texto_item)
    sobrantes = viejos[len(items):]
    for sobrante in sobrantes:
        sobrante._element.getparent().remove(sobrante._element)
    ultimo = viejos[min(len(items), len(viejos)) - 1]
    for texto_item in items[len(viejos):]:
        ultimo = clonar_despues(ultimo, texto_item)


# ===========================================================================================
# v7 - correcciones de la revision de ingenieria de software (docs/revision-tutor.md)
# ===========================================================================================

# --------------------------------------------------------------- actores contra interesados
#
# La lista mezclaba a quien opera el sistema con quien recibe su informacion, y metia
# adentro a los propios desarrolladores, que no son usuarios: eso chocaba con RF0.1, que
# dice que hay un unico par de credenciales y ninguna administracion de usuarios, y hacia
# parecer un error que los diagramas de casos de uso tengan un solo actor. El error estaba
# en la lista.

escribir(
    buscar('Con las necesidades ya definidas'),
    'Con las necesidades ya definidas, el siguiente paso fue identificar quiénes harán '
    'uso del sistema y cómo interactúan con la información. Conviene distinguir dos '
    'grupos: el actor del sistema, que es quien opera la aplicación, y los demás '
    'interesados, que reciben o aportan información pero no ingresan a ella.')

escribir(
    buscar('Sofía (Encargada del Sector)'),
    'Sofía (Encargada del Sector) — único actor del sistema. Tiene permisos totales para '
    'la carga de datos, la consulta de alertas y la generación de reportes, y es quien '
    'opera el sistema en el día a día. Es la única persona que ingresa a la aplicación, '
    'de acuerdo con RF0.1.')

escribir(
    buscar('Tamberos (Actores Operativos)'),
    'Tamberos — interesados operativos. Ejecutan las tareas diarias del establecimiento '
    '(ordeñe, vacunación, partos) y registran los datos primarios en soportes físicos '
    '—cuadernos y pizarrones— para que Sofía los cargue después. No operan el sistema.')

escribir(
    buscar('Juan Vila (Dueño)'),
    'Juan Vila (Dueño) — interesado. Recibe los reportes de producción y sanidad para la '
    'toma de decisiones económicas del establecimiento. No opera el sistema.')

escribir(
    buscar('Médico Veterinario'),
    'Médico Veterinario — interesado externo. Usa la información sanitaria y genética que '
    'el sistema produce para ajustar los planes de vacunación y las estrategias de cruce. '
    'No opera el sistema.')

# Santino y Alejo son el equipo de desarrollo, no usuarios del sistema. Su lugar es
# INTEGRANTES Y ROLES, donde ya figuran.
_desarrolladores = buscar('Administradores del Sistema')
_desarrolladores._element.getparent().remove(_desarrolladores._element)


# ------------------------------------------------------------------- objetivos verificables
#
# Cinco objetivos prometian reducir algo un ochenta, un sesenta, un cincuenta o un cuarenta
# por ciento, y nadie midio el estado anterior: sin linea base no se pueden declarar
# cumplidos ni incumplidos, y en las conclusiones hay que decir si se cumplieron. Se
# reemplazan por metas que se verifican mirando el sistema terminado.

escribir(
    buscar(BULLET + 'Centralizar la información del rodeo'),
    BULLET + 'Centralizar en un único repositorio la información del rodeo que hoy se '
    'lleva en cuadernos y pizarrones, de modo que la identificación, la producción, la '
    'reproducción y la sanidad de cada animal se consulten desde un solo lugar.')

escribir(
    buscar(BULLET + 'Disminuir los errores asociados a la transcripción'),
    BULLET + 'Reducir los errores de transcripción sustituyendo la copia manual entre '
    'soportes por un único registro validado en el momento de la carga, con los datos '
    'derivados —categoría, fecha probable de parto, fecha de secado y fin del período de '
    'descarte— calculados por el sistema y no escritos a mano.')

escribir(
    buscar(BULLET + 'Reducir en un 50% el tiempo requerido'),
    BULLET + 'Reducir el tiempo necesario para consultar la información histórica de un '
    'animal, de modo que su ficha integral —linaje, producción, reproducción y sanidad— '
    'se obtenga en una sola pantalla a partir del número de caravana.')

escribir(
    buscar(BULLET + 'Optimizar el seguimiento reproductivo'),
    BULLET + 'Dar seguimiento al ciclo reproductivo completo de cada hembra, registrando '
    'celo, servicio, tacto y parto, y manteniendo su estado reproductivo actualizado a '
    'partir de esos eventos.')

escribir(
    buscar(BULLET + 'Implementar un control de insumos'),
    BULLET + 'Implementar un control de insumos que descuente el stock al aplicarlos, '
    'avise cuando un insumo alcance su mínimo configurado y anticipe el vencimiento de '
    'cada partida, para evitar faltantes de recursos críticos.')

escribir(
    buscar(BULLET + 'Reducir en un 40% el tiempo operativo'),
    BULLET + 'Reducir la carga operativa del registro diario, de modo que cada evento se '
    'anote una sola vez y el sistema propague sus consecuencias a los demás módulos sin '
    'intervención adicional.')


# ------------------------------------------------------- requerimientos que no se verificaban

# RF1.1 decia que la categoria se ingresa y RF1.8 que el sistema la calcula. El sistema
# hace una tercera cosa: la propone y el usuario acepta o sustituye.
escribir(
    buscar('RF1.1'), 'RF1.1 Alta de animales:',
    'El sistema debe permitir registrar nuevos animales ingresando número de caravana, '
    'fecha de nacimiento, sexo y raza, y proponiendo la categoría que corresponde a partir '
    'del sexo, la fecha de nacimiento y la cantidad de partos, que el usuario puede '
    'aceptar o sustituir.')

# "Actualizar informacion" no dice que se puede cambiar: un sistema que no deja cambiar
# nada cumple el requerimiento igual. La lista sale de leer ModificarAnimal.cshtml, no de
# suponerla: la caravana tambien se modifica -se corrige una mal tipeada- y el conteo de
# partos tambien, que es como se ajusta el historial de un animal comprado.
escribir(
    buscar('RF1.3'), 'RF1.3 Modificación de animales:',
    'El sistema debe permitir modificar el número de caravana, la fecha de nacimiento, la '
    'raza, la categoría, la fotografía, la madre, el padre, la cantidad de partos '
    'registrados y, en los machos, el destino reproductivo de un animal ya dado de alta, '
    'aplicando las mismas validaciones de unicidad y de genealogía que el alta.')

# El total diario no era verificable porque RF2.3 dice que el lote y el control individual
# no se suman dentro del mismo turno: habia que decir entonces que entra en el total.
escribir(
    buscar('RF2.4'), 'RF2.4 Total diario general:',
    'El sistema debe calcular la producción total del establecimiento en una fecha como la '
    'suma de sus turnos, tomando de cada turno el registro por lote cuando existe y la '
    'suma de los controles individuales cuando el turno se anotó únicamente animal por '
    'animal, de modo que ninguna medición se cuente dos veces.')

# "Almacenar" es una afirmacion sobre la implementacion: no se puede verificar desde afuera.
escribir(
    buscar('RF2.5'), 'RF2.5 Historial productivo individual:',
    'El sistema debe registrar los controles individuales de cada animal y permitir '
    'consultarlos, imputados a la lactancia en curso en la fecha del control.')

# "Integrarse con un bot" no dice que resultado produce: un sistema que se conecta y no
# manda nada lo cumple. El comportamiento estaba en RF7.6; aca queda la configuracion.
escribir(
    buscar('RF7.5'), 'RF7.5 Configuración del canal de notificaciones:',
    'El sistema debe permitir configurar el canal de mensajería por el que envía las '
    'notificaciones, registrando el destinatario y permitiendo verificar la conexión '
    'mediante un mensaje de prueba antes de darla por activa.')


# ------------------------------------- requerimientos no funcionales que no eran medibles

# "Interfaz simple, intuitiva y de facil aprendizaje" no se puede verificar. La metrica se
# comprueba el dia de la capacitacion, y de paso da material para 2.8.
agregar_a(
    buscar('El sistema debe presentar una interfaz simple'),
    ' El criterio de verificación es que la encargada complete sin asistencia las cinco '
    'tareas de uso diario —registrar el ordeñe del turno, un celo, un servicio, un '
    'tratamiento y consultar la ficha de un animal— después de una única sesión de '
    'capacitación.')

# La mantenibilidad se apoya en una decision de arquitectura que ya se tomo, y asi se
# puede verificar leyendo el codigo.
agregar_a(
    buscar('El sistema debe ser desarrollado de tal forma'),
    ' El criterio de verificación es la separación en tres capas —presentación, dominio y '
    'persistencia—, de modo que un cambio en el acceso a datos no obligue a modificar las '
    'pantallas ni la lógica de negocio.')

# ---------------------------------------------------------------- el plan de testing
#
# El anteproyecto comprometia pruebas de caja negra y de caja blanca. Se quita la caja
# blanca: el trabajo documenta las pruebas del mismo tipo que la referencia de la
# catedra, que son funcionales, y un plan que promete algo que el proyecto despues no
# hace es peor que un plan mas corto.

escribir(
    buscar('Una vez finalizadas las funcionalidades principales'),
    'Una vez finalizadas las funcionalidades principales del proyecto se realizarán '
    'pruebas de caja negra sobre cada funcionalidad, con la finalidad de detectar '
    'errores funcionales, inconsistencias lógicas y comportamientos inesperados. Las '
    'pruebas se documentan con los datos utilizados, el resultado esperado y el '
    'obtenido.')

_titulo_blanca = buscar('Pruebas de Caja Blanca')
_cuerpo_blanca = buscar('Las pruebas de caja blanca')
for _p in (_cuerpo_blanca, _titulo_blanca):
    _p._element.getparent().remove(_p._element)

# Lo que la caja blanca aportaba -verificar los calculos en sus bordes- no se pierde:
# pasa a ser parte de lo que la caja negra comprueba, probado desde la pantalla.
agregar_a(
    buscar('El objetivo será comprobar que el sistema responda'),
    ' Se incluyen los casos de borde de los cálculos que el sistema resuelve solo —la '
    'categoría del animal en la edad de cambio, el alcance de la verificación de '
    'consanguinidad y la estimación de producción de una lactancia abierta y de una '
    'cerrada—, verificados desde la pantalla que los muestra.')


doc.save(RUTA_SALIDA)
print('guardado', RUTA_SALIDA)
